import docx

try:
    doc = docx.Document('Reliance_Carpool_Product_Document.docx')

    target_p = None
    for p in doc.paragraphs:
        if p.text.strip().startswith('6. Platform Policies'):
            target_p = p
            break
    
    if target_p:
        # Check if already added to prevent duplicates
        already_added = False
        for p in doc.paragraphs:
            if '5.9 Smart Multi-Stop Routing' in p.text:
                already_added = True
                break
        
        if not already_added:
            target_p.insert_paragraph_before('')
            h1 = target_p.insert_paragraph_before('5.9 Smart Multi-Stop Routing')
            # Let's make it look like a heading or bold
            h1.runs[0].bold = True
            target_p.insert_paragraph_before('The ride discovery system incorporates the OSRM routing engine to dynamically construct a single, optimal journey encompassing the driver’s origin, multiple accepted passenger pickup/dropoff points, and the driver’s final destination. When a new passenger joins the commute, the route and ETAs are automatically recalculated and distributed to all passengers via WebSockets in real-time.')
            
            target_p.insert_paragraph_before('')
            h2 = target_p.insert_paragraph_before('5.10 PWA Push Notifications')
            h2.runs[0].bold = True
            target_p.insert_paragraph_before('The platform implements a Service Worker architecture allowing users to subscribe to native browser Push Notifications. Even when the carpool tab is closed or minimized, passengers and hosts receive instant alerts regarding ride updates, cancellations, and SOS emergencies using secure VAPID key exchanges with the backend.')
            
            target_p.insert_paragraph_before('')
            h3 = target_p.insert_paragraph_before('5.11 Verification Queue')
            h3.runs[0].bold = True
            target_p.insert_paragraph_before('Admins have a dedicated Verification Queue dashboard to manually review and approve new Pool Hosts and their vehicle credentials before they are permitted to offer rides on the platform. This guarantees only authenticated Reliance employees and their registered vehicles provide transport.')
            
            target_p.insert_paragraph_before('')

    # Also update the Title page to say "Updated with Latest Features"
    for p in doc.paragraphs:
        if 'Version 1.0' in p.text:
            p.text = p.text.replace('Version 1.0', 'Version 1.1 (Updated with Multi-Stop & PWA)')

    doc.save('Reliance_Carpool_Product_Document.docx')
    print("SUCCESS: Original updated")

except PermissionError:
    print("PERMISSION ERROR")
except Exception as e:
    print(f"ERROR: {e}")
