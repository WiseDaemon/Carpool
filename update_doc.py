import docx
import sys

# Change stdout encoding to utf-8 if needed, but we don't need to print the doc contents anymore.
try:
    doc = docx.Document('Reliance_Carpool_Product_Document.docx')

    doc.add_heading('Key Features and Technical Capabilities', level=1)

    features = [
        "Smart Multi-Stop Routing: Dynamic route calculation utilizing the OSRM API to optimize travel paths for Pool Hosts picking up multiple passengers.",
        "PWA Push Notifications: Service worker integration for real-time offline push notifications covering ride requests, status changes, and SOS alerts.",
        "Real-time Updates: WebSocket integration (Socket.io) for live location sharing, ETAs, and immediate sync between drivers and passengers.",
        "AI Ride Assistant: Gemini-powered chatbot for booking rides, querying schedules, and fetching corporate commuting policies.",
        "Role-based Access Control: Segregated views and permissions for Pool Hosts, Passengers, and Commute Controller Admins.",
        "Incident Management & SOS: Comprehensive safety suite allowing users to trigger SOS alerts, which instantly notify Admins with location data, and automatic escalation for repeated safety flags.",
        "Penalty System: Enforced no-show logic that tracks passenger and host reliability, automatically suspending accounts after 3 infractions.",
        "Corporate SSO Ready: Mocked JWT-based authentication system ready to be linked with Reliance Active Directory.",
        "ESG Dashboard: Detailed metrics for both users and admins tracking CO2 savings, total rides, and overall platform revenue/efficiency.",
        "Women-Only Rides: Dedicated privacy feature allowing female employees to restrict visibility and matching exclusively to other women."
    ]

    for feature in features:
        p = doc.add_paragraph()
        parts = feature.split(': ', 1)
        if len(parts) == 2:
            run = p.add_run('• ' + parts[0] + ': ')
            run.bold = True
            p.add_run(parts[1])
        else:
            p.add_run('• ' + feature)

    doc.save('Reliance_Carpool_Product_Document_v2.docx')
    print("Document updated successfully.")
except Exception as e:
    print(f"Error: {e}")
